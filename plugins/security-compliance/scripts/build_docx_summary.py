#!/usr/bin/env python3
"""Build a one-page DOCX summary from a spot-check markdown report.

Pure-Python replacement for the legacy DOCX generator skill. Parses the
front matter and body of a ``commands/spot-check.md`` Step A markdown
report and emits a one-page DOCX matching the Step B contract.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Tuple

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as _docx_import_error:
    sys.stderr.write(
        "build_docx_summary: python-docx is not installed. The plugin's "
        "SessionStart hook installs it lazily; if you are running this script "
        "outside Claude Code, run `pip install --user python-docx` (or "
        "`pip install python-docx` in an activated venv) and retry.\n"
        f"build_docx_summary: import error: {_docx_import_error}\n"
    )
    sys.exit(2)

EXIT_OK = 0
EXIT_INPUT = 1
EXIT_WRITE = 2
EXIT_INTERNAL = 3

ARIAL = "Arial"
HEADER_HEX = "D5E8F0"
FOOTER_TEXT = "Confidential — for internal use"
MAX_FINDINGS = 5
MAX_ACTIONS = 3
REQUIRED_KEYS = ("date", "scenario", "applicable_frameworks", "industry")


def fail(message: str, code: int = EXIT_INPUT) -> SystemExit:
    """Print a single-line error to stderr and return a SystemExit to raise."""
    sys.stderr.write(f"build_docx_summary: {message}\n")
    return SystemExit(code)


def parse_args(argv: List[str]) -> argparse.Namespace:
    """Parse CLI args; --in and --out are required."""
    parser = argparse.ArgumentParser(
        description="Build one-page DOCX summary from a spot-check markdown report."
    )
    parser.add_argument("--in", dest="src", required=True, help="Input markdown report")
    parser.add_argument("--out", dest="dst", required=True, help="Output .docx path")
    return parser.parse_args(argv)


def split_front_matter(text: str) -> Tuple[str, str]:
    """Return ``(front_matter, body)`` or raise ``SystemExit(1)`` if malformed.

    Tolerates CRLF, UTF-8 BOM (caller passes the text from ``utf-8-sig``-decoded
    input), and a closing ``---`` at end-of-file with no trailing newline.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    if not normalized.startswith("---\n"):
        raise fail("missing opening front-matter delimiter '---'")
    end = normalized.find("\n---\n", 4)
    body_offset = 5
    if end < 0:
        if normalized.rstrip("\n").endswith("\n---"):
            end = normalized.rstrip("\n").rfind("\n---")
            body_offset = len("\n---")
        else:
            raise fail("missing closing front-matter delimiter '---'")
    return normalized[4:end], normalized[end + body_offset :]


def parse_front_matter(fm: str) -> Dict[str, Any]:
    """Mini regex-style parser for the spot-check front matter shape."""
    data: Dict[str, Any] = {}
    lines = fm.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line or line.startswith("#"):
            i += 1
            continue
        match_scalar = re.match(r"^([A-Za-z_][A-Za-z0-9_]*):\s*(.*)$", line)
        if not match_scalar:
            i += 1
            continue
        key, value = match_scalar.group(1), match_scalar.group(2).strip()
        if value == "":
            items: List[str] = []
            j = i + 1
            while j < len(lines) and lines[j].lstrip().startswith("- "):
                items.append(lines[j].lstrip()[2:].strip().strip("\"'"))
                j += 1
            data[key] = items
            i = j
            continue
        data[key] = value.strip("\"'")
        i += 1
    return data


def _check_required_keys_present(fm: Dict[str, Any]) -> None:
    """Raise ``SystemExit(1)`` when any REQUIRED_KEYS is missing."""
    for key in REQUIRED_KEYS:
        if key not in fm:
            raise fail(f"missing required front-matter key: {key}")


def _check_frameworks_list(fm: Dict[str, Any]) -> None:
    """Raise when ``applicable_frameworks`` is missing, non-list, or empty."""
    if not isinstance(fm["applicable_frameworks"], list):
        raise fail("front-matter key 'applicable_frameworks' must be a list")
    if not fm["applicable_frameworks"]:
        raise fail("front-matter key 'applicable_frameworks' must be non-empty")


def _check_required_string(fm: Dict[str, Any], key: str) -> None:
    """Raise when scalar ``key`` is not a non-empty string."""
    value = fm[key]
    if not isinstance(value, str):
        raise fail(
            f"front-matter key '{key}' must be a string scalar, got "
            f"{type(value).__name__}"
        )
    if not value.strip():
        raise fail(f"front-matter key '{key}' must be non-empty")


def _check_optional_jurisdictions(fm: Dict[str, Any]) -> None:
    """Raise when optional ``jurisdictions`` is present but not a list."""
    juris = fm.get("jurisdictions")
    if juris is not None and not isinstance(juris, list):
        raise fail("front-matter key 'jurisdictions' must be a list when present")


def validate_front_matter(fm: Dict[str, Any]) -> None:
    """Raise ``SystemExit(1)`` when required keys are missing or malformed.

    Scalar keys (``date``, ``scenario``, ``industry``) must be strings — block
    sequences in those positions parse as lists, which would otherwise pass
    a naive ``str(value).strip()`` check.
    """
    _check_required_keys_present(fm)
    _check_frameworks_list(fm)
    for required_str in ("date", "scenario", "industry"):
        _check_required_string(fm, required_str)
    _check_optional_jurisdictions(fm)


def extract_section(body: str, heading: str) -> str:
    """Return the text between a ``##`` or ``###`` heading and the next heading.

    The spot-check Step A contract says every section heading is a markdown
    ``##`` or ``###`` heading; the resolver-side validator also accepts
    levels 1-6. Match either ``##`` or ``###`` so contract-valid reports
    using ``###`` headings do not silently produce empty DOCX sections.
    """
    pattern = rf"(?ms)^#{{2,3}}\s+{re.escape(heading)}\s*$\n(.*?)(?=^#{{1,6}}\s+|\Z)"
    match = re.search(pattern, body)
    return match.group(1).strip() if match else ""


def _strip_markdown_emphasis(text: str) -> str:
    """Strip a single leading ``**bold**:`` / ``*italic*:`` prefix from an item.

    Fail-soft: if the captured group still contains a ``*`` (i.e., a second
    emphasis span follows), return the original text unchanged so the next
    bold span is not silently mangled.
    """
    match = re.match(r"^\*\*?([^*]+?)\*\*?(?:\s*[:\-]\s*)?", text)
    if not match:
        return text.strip()
    inner = match.group(1).strip()
    rest = text[match.end() :].strip()
    if "*" in rest and "*" in text[match.end() :]:
        return text.strip()
    if rest:
        return f"{inner}: {rest}".strip()
    return inner


def _extract_numbered_actions(text: str) -> List[str]:
    """Match ``N. Action`` items (matches text after ``\\d+. ``)."""
    return [
        _strip_markdown_emphasis(m.group(1))
        for m in re.finditer(r"(?m)^\s*\d+\.\s+(.+?)\s*$", text)
    ]


def _extract_bullet_actions(text: str) -> List[str]:
    """Match ``- Action`` / ``* Action`` bullet items."""
    return [
        _strip_markdown_emphasis(m.group(1))
        for m in re.finditer(r"(?m)^\s*[-*]\s+(.+?)\s*$", text)
    ]


def _extract_table_actions(text: str) -> List[str]:
    """Return first-cell text from a markdown table, skipping header + separator.

    A markdown table is ``| header | … |``, then ``|---|---|``, then data rows.
    The header row is the line immediately preceding the separator; skip both.
    """
    items: List[str] = []
    table_started = False
    for line in text.splitlines():
        stripped = line.strip()
        if not (stripped.startswith("|") and "|" in stripped[1:]):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if cells and cells[0] and set(cells[0]) <= {"-", " ", ":"}:
            table_started = True
            continue
        if not table_started:
            continue
        if cells and cells[0]:
            items.append(cells[0])
    return items


def extract_action_items(text: str, cap: int) -> List[str]:
    """Return up to ``cap`` action items from ``text``.

    Step A's "Recommended Next Steps" contract is a prioritized list, but the
    template does not constrain the renderer to numbered items vs bullets vs
    table rows. Try all three shapes in priority order; first non-empty wins:

    1. Numbered list  ``1. Action.``
    2. Unordered list ``- Action.`` / ``* Action.``
    3. Table row      data rows after a ``|---|---|`` separator
    """
    items = (
        _extract_numbered_actions(text)
        or _extract_bullet_actions(text)
        or _extract_table_actions(text)
    )
    return items[:cap]


def extract_findings_sentences(text: str, cap: int) -> List[str]:
    """Return up to ``cap`` finding sentences (paragraph + sentence split fallback)."""
    if not text.strip():
        return []
    flat = re.sub(r"\s+", " ", text.strip())
    sentences = re.split(r"(?<=[.!?])\s+", flat)
    return [s for s in sentences if s][:cap]


def extract_tldr_bullets(body: str, cap: int) -> List[str]:
    """Return up to ``cap`` bullets from the TL;DR section (fallback for findings)."""
    section = extract_section(body, "TL;DR")
    if not section:
        return []
    bullets = re.findall(r"(?m)^\s*[-*]\s+(.+?)\s*$", section)
    return bullets[:cap]


def collect_findings(body: str, min_count: int = 3, max_count: int = 5) -> List[str]:
    """Pull findings from ``Compliance Gap Summary``; fall back to TL;DR bullets.

    Step B requires 3-5 findings. When the gap-summary section yields fewer
    than ``min_count`` sentences, top up from the TL;DR bullet list.
    """
    primary = extract_findings_sentences(
        extract_section(body, "Compliance Gap Summary"), max_count
    )
    if len(primary) >= min_count:
        return primary[:max_count]
    seen = {item.lower(): None for item in primary}
    for bullet in extract_tldr_bullets(body, max_count):
        if bullet.lower() in seen:
            continue
        primary.append(bullet)
        seen[bullet.lower()] = None
        if len(primary) >= max_count:
            break
    return primary[:max_count]


def display_framework(name: str) -> str:
    """Render canonical SOC 2 / ISO 27001 / etc. display names."""
    mapping = {
        "SOC2": "SOC 2",
        "ISO_27001": "ISO 27001",
        "PCI_DSS": "PCI DSS",
        "EU_AI_ACT": "EU AI Act",
        "CCPA_CPRA": "CCPA/CPRA",
        "NIST_CSF": "NIST CSF",
    }
    return mapping.get(name, name)


def shade_cell(cell: Any, fill_hex: str) -> None:
    """Apply solid background fill to a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def configure_page(document: Any) -> None:
    """Set US Letter, 1-inch margins, default Arial 10pt body."""
    section = document.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    style = document.styles["Normal"]
    style.font.name = ARIAL
    style.font.size = Pt(10)


def add_header(document: Any, fm: Dict[str, Any]) -> None:
    """Add the centered bold 16pt title plus the metadata sub-line."""
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Compliance Assessment Summary")
    title_run.bold = True
    title_run.font.name = ARIAL
    title_run.font.size = Pt(16)

    juris = ", ".join(fm.get("jurisdictions", [])) or "—"
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Date: {fm['date']}    Industry: {fm['industry']}    Jurisdictions: {juris}"
    )
    meta_run.font.name = ARIAL
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_section_heading(document: Any, text: str) -> None:
    """Add a bold 12pt section header."""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = ARIAL
    run.font.size = Pt(12)


def add_scenario(document: Any, scenario: str) -> None:
    """Add the Scenario block (2-3 sentences)."""
    add_section_heading(document, "Scenario")
    para = document.add_paragraph()
    para.add_run(scenario).font.name = ARIAL


def add_frameworks_table(document: Any, frameworks: List[str]) -> None:
    """Add the Applicable Frameworks compact table with shaded header row."""
    add_section_heading(document, "Applicable Frameworks")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Framework", "Status", "Key Obligation")
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = ARIAL
        shade_cell(cell, HEADER_HEX)
    for fw in frameworks:
        row = table.add_row().cells
        row[0].text = display_framework(fw)
        row[1].text = "Applicable"
        row[2].text = "See full report"


def add_findings(document: Any, findings: List[str]) -> None:
    """Add the Key Findings numbered list (capped to MAX_FINDINGS)."""
    add_section_heading(document, "Key Findings")
    for idx, item in enumerate(findings, start=1):
        para = document.add_paragraph()
        para.add_run(f"{idx}. {item}").font.name = ARIAL


def add_actions(document: Any, actions: List[str]) -> None:
    """Add the Priority Actions numbered list (capped to MAX_ACTIONS)."""
    add_section_heading(document, "Priority Actions")
    for idx, item in enumerate(actions, start=1):
        para = document.add_paragraph()
        para.add_run(f"{idx}. {item}").font.name = ARIAL


def add_footer(document: Any) -> None:
    """Add the centered 8pt-gray footer with the page-number field."""
    footer = document.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{FOOTER_TEXT}    Page ")
    run.font.name = ARIAL
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = para.add_run()
    page_run.font.name = ARIAL
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)


def build(src: Path, dst: Path) -> int:
    """Read ``src`` markdown, build the DOCX, write to ``dst``. Returns exit code.

    Reads with ``utf-8-sig`` so a leading BOM (e.g., Windows-edited reports)
    is silently consumed. ``UnicodeError`` is reported as an input error, not
    a generic write failure.
    """
    try:
        text = src.read_text(encoding="utf-8-sig")
    except OSError as exc:
        sys.stderr.write(f"build_docx_summary: cannot read input: {exc}\n")
        return EXIT_INPUT
    except UnicodeError as exc:
        sys.stderr.write(f"build_docx_summary: cannot decode input as UTF-8: {exc}\n")
        return EXIT_INPUT

    fm_raw, body = split_front_matter(text)
    fm = parse_front_matter(fm_raw)
    validate_front_matter(fm)

    scenario = str(fm["scenario"]).strip()
    findings = collect_findings(body, min_count=3, max_count=MAX_FINDINGS)
    actions_text = extract_section(body, "Recommended Next Steps")
    actions = extract_action_items(actions_text, MAX_ACTIONS)

    try:
        document = Document()
        configure_page(document)
        add_header(document, fm)
        add_scenario(document, scenario)
        add_frameworks_table(document, fm["applicable_frameworks"])
        add_findings(document, findings)
        add_actions(document, actions)
        add_footer(document)
    except ValueError as exc:
        sys.stderr.write(
            "build_docx_summary: input contains characters incompatible with "
            f"DOCX XML: {exc}\n"
        )
        return EXIT_INPUT

    try:
        dst.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(dst))
    except OSError as exc:
        sys.stderr.write(f"build_docx_summary: cannot write output: {exc}\n")
        return EXIT_WRITE
    except ValueError as exc:
        sys.stderr.write(
            "build_docx_summary: serialized DOCX rejected by writer "
            f"(XML-incompatible content): {exc}\n"
        )
        return EXIT_INPUT
    return EXIT_OK


def main(argv: List[str]) -> int:
    """CLI entry point. Returns exit code."""
    args = parse_args(argv)
    return build(Path(args.src), Path(args.dst))


if __name__ == "__main__":
    try:
        sys.exit(main(sys.argv[1:]))
    except SystemExit:
        raise
    except Exception as exc:
        sys.stderr.write(f"build_docx_summary: unexpected error: {exc}\n")
        sys.exit(EXIT_INTERNAL)
