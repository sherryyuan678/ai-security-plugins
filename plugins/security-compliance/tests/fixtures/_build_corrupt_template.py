"""Build tests/fixtures/corrupt_template.docx — valid zip; truncated XML.

Used by failure-mode tests to assert EXIT_INTERNAL=3 on validate failure.

Corrupts BOTH word/document.xml AND word/styles.xml — apply-template grafts
styles.xml from template, so corrupting it surfaces in the grafted output's
xml_wellformed check (which is what the test asserts). Corrupting just
word/document.xml would NOT trigger failure because apply-template keeps
the target's document.xml intact (only styles/theme/header/footer/media
are grafted from the template).

Run: python3 tests/fixtures/_build_corrupt_template.py
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document


def main() -> None:
    here = Path(__file__).resolve().parent
    out = here / "corrupt_template.docx"

    doc = Document()
    doc.add_heading("Corrupt Template (will be truncated)", level=1)
    doc.add_paragraph("This document is intentionally invalid — for testing.")
    doc.save(str(out))

    src_bytes = out.read_bytes()
    src = zipfile.ZipFile(io.BytesIO(src_bytes), "r")
    out_buf = io.BytesIO()
    corrupted_count = 0
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as outz:
        for name in src.namelist():
            data = src.read(name)
            if name in ("word/document.xml", "word/styles.xml"):
                txt = data.decode("utf-8")
                for marker in ("<w:p ", "<w:p>", "<w:body>", "<w:style ", "<w:styles "):
                    cut = txt.find(marker)
                    if cut > 0:
                        txt = txt[: cut + len(marker)] + "<<<INVALID_PARSE"
                        data = txt.encode("utf-8")
                        corrupted_count += 1
                        break
            outz.writestr(name, data)

    if corrupted_count == 0:
        raise RuntimeError("could not find any opener to corrupt")

    out.write_bytes(out_buf.getvalue())
    print(f"wrote {out} ({out.stat().st_size} bytes); corrupted {corrupted_count} part(s)")
    print("note: word/document.xml AND word/styles.xml are intentionally malformed")


if __name__ == "__main__":
    main()
