"""Build tests/fixtures/reviewer_marked.docx — text with <w:ins> and <w:del>.

The python-docx API does not expose tracked-change construction directly, so
we build a normal docx and then post-process word/document.xml to wrap a few
runs in <w:ins> and <w:del>. The result has 2 inserts and 1 delete that
soffice's --accept-changes will resolve.

Run: python3 tests/fixtures/_build_reviewer_marked.py
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document


def _patch_document_xml(xml_text: str) -> str:
    """Replace specific runs with <w:ins> and <w:del> wrappers."""
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if f'xmlns:w="{NS_W}"' not in xml_text:
        # Document already has w: prefix bound; we don't touch it.
        pass

    # Insertion: wrap the placeholder phrase "INS_TARGET" in <w:ins>
    xml_text = xml_text.replace(
        '<w:r><w:t>INS_TARGET</w:t></w:r>',
        '<w:ins w:id="1" w:author="Reviewer" w:date="2026-05-06T00:00:00Z">'
        '<w:r><w:t>INS_TARGET</w:t></w:r></w:ins>',
    )
    # Deletion: wrap "DEL_TARGET"
    xml_text = xml_text.replace(
        '<w:r><w:t>DEL_TARGET</w:t></w:r>',
        '<w:del w:id="2" w:author="Reviewer" w:date="2026-05-06T00:00:00Z">'
        '<w:r><w:delText>DEL_TARGET</w:delText></w:r></w:del>',
    )
    return xml_text


def main() -> None:
    here = Path(__file__).resolve().parent
    out = here / "reviewer_marked.docx"

    doc = Document()
    doc.add_heading("Reviewer-Marked Document", level=1)
    # Build paragraph with one run per phrase so each phrase has its own
    # <w:r><w:t>...</w:t></w:r> wrapper that we can target.
    p = doc.add_paragraph()
    p.add_run("Original text. ")
    p.add_run("INS_TARGET")
    p.add_run(" middle ")
    p.add_run("DEL_TARGET")
    p.add_run(" End text.")
    doc.save(str(out))

    # Post-process: re-zip, patching word/document.xml with tracked-change wrappers
    src_bytes = out.read_bytes()
    src = zipfile.ZipFile(io.BytesIO(src_bytes), "r")
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as outz:
        for name in src.namelist():
            data = src.read(name)
            if name == "word/document.xml":
                txt = data.decode("utf-8")
                txt = _patch_document_xml(txt)
                data = txt.encode("utf-8")
            outz.writestr(name, data)
    out.write_bytes(out_buf.getvalue())
    print(f"wrote {out} ({out.stat().st_size} bytes)")

    # Verify the tracked-change tags are present
    final = zipfile.ZipFile(out, "r")
    body = final.read("word/document.xml").decode("utf-8")
    assert "<w:ins" in body, "expected <w:ins> in output"
    assert "<w:del" in body, "expected <w:del> in output"
    print("verified: <w:ins> and <w:del> present")


if __name__ == "__main__":
    main()
