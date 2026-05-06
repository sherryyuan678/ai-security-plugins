"""Build tests/fixtures/acme_template.docx — theme + header + footer + media.

Differs from plain_summary in: theme color, header text, footer text, custom
paragraph style, and embedded PNG image (so word/media/* exists).

Run: python3 tests/fixtures/_build_acme_template.py
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor


def _minimal_png_bytes() -> bytes:
    """Return bytes for a 1x1 transparent PNG (smallest valid).

    Anatomy: 8-byte signature + IHDR (13B data + 4B CRC) + IDAT + IEND.
    """
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(tag: bytes, data: bytes) -> bytes:
        length = struct.pack(">I", len(data))
        crc = struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        return length + tag + data + crc

    # IHDR: width=1, height=1, bit_depth=8, color_type=6 (RGBA), filter=0, compression=0, interlace=0
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
    # IDAT: zlib-compressed scanline (1 px RGBA, 1 filter byte = 5 bytes total)
    raw = b"\x00\x00\x00\x00\x00"
    idat = zlib.compress(raw, 9)
    iend = b""

    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", iend)


def main() -> None:
    here = Path(__file__).resolve().parent
    out = here / "acme_template.docx"

    doc = Document()

    # Custom style
    styles = doc.styles
    custom = styles.add_style("AcmeAccent", WD_STYLE_TYPE.PARAGRAPH)
    custom.font.name = "Arial"
    custom.font.size = Pt(13)
    custom.font.color.rgb = RGBColor(0x18, 0x4F, 0x9E)

    # Add header + footer text
    sect = doc.sections[0]
    sect.header.paragraphs[0].text = "ACME Corp — CONFIDENTIAL"
    sect.footer.paragraphs[0].text = "ACME Corp © 2026 — Page footer"

    doc.add_heading("ACME Template Document", level=1)
    doc.add_paragraph(
        "This is the styling template. After --apply-template, the target's "
        "theme, headers, footers, and media should be replaced by these.",
        style="AcmeAccent",
    )

    # Add an inline image so word/media/ has content
    img_bytes = _minimal_png_bytes()
    img_stream = io.BytesIO(img_bytes)
    doc.add_picture(img_stream, width=Inches(0.5))

    doc.add_paragraph("End of template body.")

    doc.save(str(out))
    print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
