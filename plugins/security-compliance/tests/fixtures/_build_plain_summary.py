"""Build tests/fixtures/plain_summary.docx — heading + 3 paras + 1 table.

Run: python3 tests/fixtures/_build_plain_summary.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def main() -> None:
    here = Path(__file__).resolve().parent
    out = here / "plain_summary.docx"

    doc = Document()
    doc.add_heading("Quarterly Compliance Summary", level=1)
    doc.add_paragraph(
        "This document is a synthetic test fixture for the security-compliance "
        "plugin's edit-docx pipeline. It exercises the apply-template path."
    )
    doc.add_paragraph(
        "The body holds three paragraphs and one small table. Theme, header, "
        "footer, and styles are the python-docx defaults."
    )
    doc.add_paragraph(
        "After --apply-template, the styles, theme, headers, footers, and media "
        "should match the template's; numbering should remain target's."
    )

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Control"
    hdr[1].text = "Framework"
    hdr[2].text = "Status"
    row = table.rows[1].cells
    row[0].text = "CC6.1"
    row[1].text = "SOC 2"
    row[2].text = "Compliant"

    for sect in doc.sections:
        sect.page_height = Inches(11)
        sect.page_width = Inches(8.5)

    for p in doc.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)

    doc.save(str(out))
    print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
